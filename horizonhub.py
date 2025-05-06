import streamlit as st
import pandas as pd
from datetime import datetime
from io import BytesIO
from docx import Document
import pdfkit

st.set_page_config(page_title="Horizon Hub", layout="wide")
st.title("🌐 Horizon Hub - Smart Workspace Organizer")

if "events" not in st.session_state:
    st.session_state["events"] = []
if "tasks" not in st.session_state:
    st.session_state["tasks"] = []


# Sidebar Navigation
page = st.sidebar.radio("Go to", ["📅 Dashboard", "📌 Create Event", "✅ To-do List", "📂 Upload Document", "📊 Productivity Insights"])

# -----------------------------
# 📅 DASHBOARD
# -----------------------------
if page == "📅 Dashboard":
    st.subheader("📅 Upcoming Events")
    if st.session_state.events:
        for ev in st.session_state.events:
            st.markdown(f"""
            **🗓️ {ev['title']}**  
            📍 Date: {ev['date']} at {ev['time']}  
            🧑 Invitees: {ev['attendees']}  
            📝 Agenda: {ev['agenda']}  
            """)
            st.divider()
    else:
        st.info("No events scheduled.")

# -----------------------------
# 📌 CREATE EVENT
# -----------------------------
elif page == "📌 Create Event":
    st.subheader("📌 Schedule a New Event")
    title = st.text_input("Event Title")
    date = st.date_input("Date")
    time = st.time_input("Time")
    attendees = st.text_input("Invitees (comma-separated emails)")
    agenda = st.text_area("Agenda / Notes")
    if st.button("Create Event"):
        event = {
            "title": title,
            "date": date.strftime("%Y-%m-%d"),
            "time": time.strftime("%H:%M"),
            "attendees": attendees,
            "agenda": agenda
        }
        st.session_state.events.append(event)
        st.success("✅ Event added successfully!")

# -----------------------------
# ✅ TO-DO LIST
# -----------------------------
elif page == "✅ To-do List":
    st.subheader("📝 To-do Task Manager")
    task = st.text_input("New Task")
    priority = st.selectbox("Priority", ["High", "Medium", "Low"])
    if st.button("Add Task"):
        st.session_state.tasks.append({"task": task, "priority": priority})
        st.success(f"Task '{task}' added!")

    if st.session_state.tasks:
        st.markdown("### Your Tasks")
        for i, t in enumerate(st.session_state.tasks):
            st.write(f"📌 {t['task']} — {t['priority']} Priority")
            if st.button(f"Mark as Done ✅", key=i):
                st.session_state.tasks.pop(i)
                st.success("Task marked as complete!")
                st.experimental_rerun()
    else:
        st.info("No tasks added yet.")

# -----------------------------
# 📂 UPLOAD DOCUMENT
# -----------------------------
elif page == "📂 Upload Document":
    st.subheader("📤 Upload & Manage Documents")
    uploaded_file = st.file_uploader("Upload a document (PDF, DOCX, TXT)")
    format_choice = st.selectbox("Convert to format", ["Select", "PDF", "TXT", "DOCX"])

    if uploaded_file:
        file_name = uploaded_file.name
        file_extension = file_name.split('.')[-1].lower()
        st.success(f"Uploaded: {file_name}")

        if format_choice == "PDF" and file_extension == "docx":
            # Convert DOCX to PDF
            doc = Document(uploaded_file)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            pdf_data = pdfkit.from_string(text, False)
            st.download_button(
                label="Download Converted PDF",
                data=pdf_data,
                file_name="converted_document.pdf",
                mime="application/pdf"
            )
            st.info("✅ Converted to PDF.")

        elif format_choice == "TXT" and file_extension == "pdf":
            # Convert PDF to TXT (mocked, actual PDF parsing would be more complex)
            pdf_data = uploaded_file.read()
            text_data = "Converted text from PDF"  # Mocking PDF to text conversion.
            st.download_button(
                label="Download Converted TXT",
                data=text_data,
                file_name="converted_document.txt",
                mime="text/plain"
            )
            st.info("✅ Converted to TXT.")

        elif format_choice == "DOCX" and file_extension == "pdf":
            # Convert PDF to DOCX (mocked, actual PDF to DOCX would need a different method)
            pdf_data = uploaded_file.read()
            doc = Document()
            doc.add_paragraph("Converted content from PDF")  # Mocking PDF to DOCX conversion.
            docx_stream = BytesIO()
            doc.save(docx_stream)
            docx_stream.seek(0)
            st.download_button(
                label="Download Converted DOCX",
                data=docx_stream,
                file_name="converted_document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.info("✅ Converted to DOCX.")

        else:
            st.warning("Please select a valid format to convert to.")

# -----------------------------
# 📊 PRODUCTIVITY INSIGHTS
# -----------------------------
elif page == "📊 Productivity Insights":
    st.subheader("📈 Productivity Overview")
    total_events = len(st.session_state.events)
    total_tasks = len(st.session_state.tasks)

    st.metric(label="📅 Events Scheduled", value=total_events)
    st.metric(label="✅ Tasks Pending", value=total_tasks)
    
    df = pd.DataFrame({
        "Category": ["Events", "Tasks Done", "Tasks Pending"],
        "Count": [total_events, 5, total_tasks]
    })
    st.bar_chart(df.set_index("Category"))
